from tkinter import *
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import datetime
import os
import mysql.connector

mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    passwd="Mehul98@",
    database="invoice"
)

class Invoice:

    total = 0
    index = 0
    item_list = {}
    total_row = 8
    message_string = "SRNO\t\tITEM\t\tQUANTITY\tRATE\t\tPRICE"
    mycursor = mydb.cursor()
    prev_billno = 0


    def __init__(self, root):
        root.title("Invoice")

        self.label_0 = Label(root, text="HMR ARTS", width=50, font=("bold", 20)).grid(column=0, row=0, columnspan=3)

        self.label_from = Label(root, text="FROM:").grid(row=2, column=0, sticky="E")
        self.entry_from = Entry(root, bd=5)
        self.entry_from.grid(row=2, column=1, sticky="W", pady=5)

        self.label_to = Label(root, text="TO:").grid(row=3, column=0, sticky="E")
        self.entry_to = Entry(root, bd=5)
        self.entry_to.grid(row=3, column=1, sticky="W", pady=5)

        self.label_cntitems = Label(root, text="No of ITEMS:").grid(row=4, column=0, sticky="E")
        self.entry_cntitems = Entry(root, bd=5)
        self.entry_cntitems.grid(row=4, column=1, sticky="W", pady=5)

        self.label_item = Label(root, text="ITEMS").grid(row=5, column=0, sticky="W")
        self.label_quantity = Label(root, text="QUANTITY").grid(row=5, column=1, sticky="W")
        self.label_rate = Label(root, text="RATE").grid(row=5, column=2, sticky="W")
        e1 = self.entry_item = Entry(root, bd=5)
        e1.grid(row=6, column=0, sticky="W")
        e2 = self.entry_quantity = Entry(root, bd=5)
        e2.grid(row=6, column=1, sticky="W")
        e3 = self.entry_rate = Entry(root, bd=5)
        e3.grid(row=6, column=2, sticky="W")

        self.label_message = Label(root)
        self.additembutton = Button(root, text="ADD ITEM", command=self.item_collection).grid(row=7, column=2,
                                                                                              sticky="W")

        self.items_label = Label(root)

        self.label_total = Label(root, text="TOTAL:")

        self.label_anstotal = Label(root, bd=5)

        self.save_button = Button(root, text="SAVE", command=self.make_word)

    def item_collection(self):
        self.index += 1
        Invoice.total_row += 1
        curr_tcount=self.counttotal()

        Invoice.item_list[self.entry_item.get()] = [self.entry_quantity.get(), self.entry_rate.get(), curr_tcount]

        keys = self.entry_item.get()
        new_str = str(self.index) + "\t\t" + str(keys) + "\t\t" + str(self.item_list[keys][0]) + "\t\t" + str(self.item_list[keys][1] + "\t\t" + str(self.item_list[keys][2]))

        self.message_string = self.message_string+"\n"+new_str

        self.items_label.config(text=self.message_string)
        self.items_label.grid(row=self.total_row - 1, column=0, columnspan=2)


        # success mess label
        self.label_message.config(text="Successfully Added")
        self.label_message.grid(row=7, column=0)

        # total label display
        self.label_total.grid(row=self.total_row, column=0, sticky="E")

        Invoice.total = Invoice.total + curr_tcount

        self.entry_item.delete(0, END)
        self.entry_quantity.delete(0, END)
        self.entry_rate.delete(0, END)

        self.label_anstotal.config(text=str(self.total))
        self.label_anstotal.grid(row=self.total_row, column=1, sticky="W", pady=5)

        self.save_button.grid(row=self.total_row+1, column=2, sticky="W")

    def counttotal(self):
        rate = self.entry_rate.get()
        quantity = self.entry_quantity.get()
        return int(rate) * int(quantity)

    def make_word(self):

        cnt = 0
        Invoice.prev_billno = int(self.get_billno())

        document = Document()
        document.add_paragraph("Bill No:" + str(self.prev_billno+1))

        x = datetime.datetime.now()
        para0 = document.add_paragraph(str(x.strftime("%x")))
        para0.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        style = document.styles['Normal']
        font = style.font
        font.size = Pt(18)

        head = document.add_paragraph('H.M.R Art')
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER

        para1 = document.add_paragraph('Name: ')
        run = para1.add_run("Ashwin Gohil")
        run.bold = True
        para2 = document.add_paragraph("To: ")
        run = para2.add_run(str(self.entry_to.get()))
        run.bold = True

        table = document.add_table(rows=1, cols=5)

        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'SR NO'
        heading_cells[1].text = 'ITEM'
        heading_cells[2].text = 'QUANTITY'
        heading_cells[3].text = 'RATE'
        heading_cells[4].text = 'PRICE'

        for keys in self.item_list:
            cnt += 1
            cells = table.add_row().cells
            cells[0].text = str(cnt)
            cells[1].text = str(keys)
            cells[2].text = str(self.item_list[keys][0])
            cells[3].text = str(self.item_list[keys][1])
            cells[4].text = str(self.item_list[keys][2])

        table.style = 'LightShading-Accent1'

        para3 = document.add_paragraph("Total Amount: ")
        run = para3.add_run(str(self.total))
        run.bold = True
        para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para3.space_before = Pt(18)

        sign = document.add_paragraph('Receivers Sign: _________________')
        sign1 = document.add_paragraph('Authorised Sign: _________________')
        sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.save('test.docx')

        self.populate_db()

        # clearing from to and no of items entries
        self.entry_from.delete(0, END)
        self.entry_to.delete(0, END)
        self.entry_cntitems.delete(0, END)

        fpath = 'C:\\Users\\MAG\\Desktop\\Invoicepro\\test.docx'
        os.startfile(fpath, 'open')

        Invoice.item_list.clear()
        Invoice.total = 0
        self.label_message.config(text="")
        self.items_label.config(text="")
        self.message_string = "SRNO\t\tITEM\t\tQUANTITY\tRATE\t\tPRICE"
        self.label_anstotal.config(text="")
        self.index = 0

    def populate_db(self):

        sqlFormula = "INSERT INTO bill_detail (RECEIVERS_NAME,TOTAL_AMT) VALUES(%s, %s)"
        sqlFormula1 = "INSERT INTO bill_items_detail (BILL_NO, ITEM_NAME, RATE, QUANTITY) VALUES(%s, %s, %s, %s)"

        val = [str(self.entry_to.get()), str(self.total)]
        self.mycursor.execute(sqlFormula, val)

        for keys in self.item_list:
            val = [self.prev_billno+1, str(keys), int(self.item_list[keys][0]), int(self.item_list[keys][1])]
            self.mycursor.execute(sqlFormula1, val)

        mydb.commit()

    def get_billno(self):

        sql_formula = "Select BILL_NO from bill_detail Order By BILL_NO Desc Limit 1"
        self.mycursor.execute(sql_formula)

        val = self.mycursor.fetchall()

        return val[0][0]

master = Tk()
i = Invoice(master)
master.mainloop()
